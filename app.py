import streamlit as st
import os
import time

from docx import Document
from deep_translator import GoogleTranslator
import pdfplumber
from docx2pdf import convert

# ---------------- CONFIG ----------------
st.set_page_config(
    page_title="International Document Translator",
    layout="wide"
)

# ---------------- CONSTANTS ----------------
LANGUAGES = {
    "English 🇬🇧": "en",
    "Chinese 🇨🇳": "zh-CN",
    "Vietnamese 🇻🇳": "vi",
    "Japanese 🇯🇵": "ja"
}

TEMP_DIR = "temp"
os.makedirs(TEMP_DIR, exist_ok=True)

# ---------------- HEADER ----------------
st.markdown(
    """
    <h1 style="text-align:center;">
        🌍 International Academic Document Translator
    </h1>
    <p style="text-align:center; font-size:18px;">
        APS • German Universities • International Students
    </p>
    """,
    unsafe_allow_html=True
)

st.divider()

# ---------------- SIDEBAR ----------------
with st.sidebar:
    st.header("⚙️ Translation Settings")

    target_language_label = st.selectbox(
        "Select Target Language",
        list(LANGUAGES.keys())
    )

    target_lang = LANGUAGES[target_language_label]

    st.markdown("---")

    st.markdown("### 📌 Supported Formats")
    st.markdown("• Word (.docx) – **Best Quality**")
    st.markdown("• PDF (.pdf) – **Best Effort**")

    st.markdown("---")

    st.info(
        "For perfect layout preservation, "
        "upload Word documents whenever possible."
    )

# ---------------- FILE UPLOAD ----------------
st.subheader("📤 Upload Your Academic Document")

uploaded_file = st.file_uploader(
    "Upload Word or PDF file",
    type=["docx", "pdf"]
)

if not uploaded_file:
    st.stop()

file_name = uploaded_file.name
file_ext = file_name.split(".")[-1].lower()

input_path = os.path.join(TEMP_DIR, f"input.{file_ext}")
translated_docx = os.path.join(TEMP_DIR, "translated.docx")
translated_pdf = os.path.join(TEMP_DIR, "translated.pdf")

with open(input_path, "wb") as f:
    f.write(uploaded_file.read())
from deep_translator import GoogleTranslator

sample_text = ""
if file_ext == "docx":
    from docx import Document
    d = Document(input_path)
    if d.paragraphs:
        sample_text = d.paragraphs[0].text
elif file_ext == "pdf":
    import pdfplumber
    with pdfplumber.open(input_path) as pdf:
        if pdf.pages:
            sample_text = pdf.pages[0].extract_text()

if sample_text:
    detector = GoogleTranslator(source="auto", target=target_lang)
    detected_lang = detector.detect(sample_text)
    st.info(f"Detected source language: {detected_lang}")

# ---------------- FILE ANALYSIS ----------------
st.divider()
st.subheader("🔍 Document Analysis")

colA, colB, colC = st.columns(3)

with colA:
    st.metric("File Name", file_name)

with colB:
    st.metric("File Type", file_ext.upper())

with colC:
    st.metric("Target Language", target_language_label)

if file_ext == "pdf":
    st.warning(
        "PDF detected. Layout may slightly change after translation. "
        "This is a known limitation of PDF technology."
    )

# ---------------- TRANSLATION PIPELINE ----------------
st.divider()
st.subheader("🔄 Translation Pipeline")

progress = st.progress(0)
status = st.empty()

def translate_text(text):
    if not text or not text.strip():
        return text
    try:
        translator = GoogleTranslator(source="de", target=target_lang)
        time.sleep(0.6)
        return translator.translate(text)
    except:
        return text

# ---------------- WORD PIPELINE ----------------
if file_ext == "docx":
    status.info("Reading Word document...")
    progress.progress(10)

    doc = Document(input_path)
    new_doc = Document()

    status.info("Translating paragraphs...")
    progress.progress(30)

    for para in doc.paragraphs:
        translated = translate_text(para.text)
        new_doc.add_paragraph(translated)

    status.info("Translating tables...")
    progress.progress(60)

    for table in doc.tables:
        new_table = new_doc.add_table(
            rows=len(table.rows),
            cols=len(table.columns)
        )
        for r, row in enumerate(table.rows):
            for c, cell in enumerate(row.cells):
                new_table.rows[r].cells[c].text = translate_text(cell.text)

    new_doc.save(translated_docx)

# ---------------- PDF PIPELINE ----------------
else:
    status.info("Reading PDF document...")
    progress.progress(10)

    new_doc = Document()

    with pdfplumber.open(input_path) as pdf:
        total_pages = len(pdf.pages)

        for i, page in enumerate(pdf.pages):
            status.info(f"Translating page {i+1} / {total_pages}")
            progress.progress(10 + int((i / total_pages) * 60))

            text = page.extract_text()
            if text:
                translated = translate_text(text)
                new_doc.add_paragraph(translated)

    new_doc.save(translated_docx)

# ---------------- PDF EXPORT ----------------
status.info("Generating translated PDF...")
progress.progress(90)

convert(translated_docx, translated_pdf)

progress.progress(100)
status.success("Translation completed successfully!")

# ---------------- RESULTS ----------------
st.divider()
st.subheader("📥 Download Translated Files")

col1, col2 = st.columns(2)

with col1:
    with open(translated_docx, "rb") as f:
        st.download_button(
            "⬇ Download Translated Word",
            f,
            file_name="translated.docx"
        )

with col2:
    with open(translated_pdf, "rb") as f:
        st.download_button(
            "⬇ Download Translated PDF",
            f,
            file_name="translated.pdf"
        )

# ---------------- FOOTER ----------------
st.divider()
st.markdown(
    """
    <p style="text-align:center; color:gray;">
        Built for international students • APS & German universities
    </p>
    """,
    unsafe_allow_html=True
)
