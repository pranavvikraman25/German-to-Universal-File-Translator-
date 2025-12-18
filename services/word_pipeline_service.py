from docx import Document
from services.translation_service import translate_block

def process_word(input_path, output_path, target_lang):
    src_doc = Document(input_path)
    out_doc = Document()

    for para in src_doc.paragraphs:
        out_doc.add_paragraph(
            translate_block(para.text, target_lang)
        )

    for table in src_doc.tables:
        new_table = out_doc.add_table(
            rows=len(table.rows),
            cols=len(table.columns)
        )

        for r, row in enumerate(table.rows):
            for c, cell in enumerate(row.cells):
                new_table.rows[r].cells[c].text = translate_block(
                    cell.text,
                    target_lang
                )

    out_doc.save(output_path)
