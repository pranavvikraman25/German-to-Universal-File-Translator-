from docx import Document
from utils.translator_utils import translate_text

def translate_word(input_path, output_path, target_lang):
    doc = Document(input_path)
    new_doc = Document()

    for para in doc.paragraphs:
        new_doc.add_paragraph(
            translate_text(para.text, target_lang)
        )

    for table in doc.tables:
        new_table = new_doc.add_table(
            rows=len(table.rows),
            cols=len(table.columns)
        )

        for r, row in enumerate(table.rows):
            for c, cell in enumerate(row.cells):
                new_table.rows[r].cells[c].text = translate_text(
                    cell.text,
                    target_lang
                )

    new_doc.save(output_path)
